#!/usr/bin/env python3
"""
Build DOCX and PDF reports from medsyn_dialogue_analysis outputs.
"""

from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib import colors


def add_table_from_df_docx(doc: Document, df: pd.DataFrame, title: str, floatfmt: str = "{:.3f}") -> None:
    doc.add_paragraph(title).runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, v in enumerate(row):
            if isinstance(v, float):
                cells[j].text = floatfmt.format(v)
            else:
                cells[j].text = str(v)
    doc.add_paragraph("")


def df_to_rl_table(df: pd.DataFrame):
    data = [list(df.columns)] + df.astype(str).values.tolist()
    tbl = Table(data)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return tbl


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--out_dir", required=True, type=Path, help="Output directory from medsyn_dialogue_analysis.py")
    args = ap.parse_args()

    out_dir: Path = args.out_dir
    fig_dir = out_dir / "figures"

    case_df = pd.read_csv(out_dir / "case_level.csv")
    turn_df = pd.read_csv(out_dir / "turn_level.csv")
    q_cat = pd.read_csv(out_dir / "question_category_distribution.csv")

    # Summary tables
    case_summary = (
        case_df.groupby(["session", "group"])
        .agg(n_cases=("note_id", "count"), mean_turns=("turn_pairs", "mean"), median_turns=("turn_pairs", "median"), mean_duration_min=("duration_sec", lambda x: x.mean() / 60))
        .reset_index()
    )
    turn_summary = (
        turn_df.groupby(["session", "group"])
        .agg(n_turns=("turn_index", "count"), mean_overlap_ctx=("overlap_context", "mean"), prop_has_note=("a_has_note", "mean"), prop_not_in_note=("a_not_in_note", "mean"), mean_specificity=("a_specificity", "mean"))
        .reset_index()
    )
    q_tbl = q_cat.pivot_table(index=["session", "group"], columns="q_category", values="prop").fillna(0).reset_index()

    # DOCX
    doc = Document()
    title = doc.add_paragraph("MedSyn Interactive Dialogue Analysis (Session 2 & Session 4)")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_heading("Quantitative summaries", level=1)
    add_table_from_df_docx(doc, case_summary.round(3), "Table 1. Case-level summary.", floatfmt="{:.3f}")
    add_table_from_df_docx(doc, q_tbl.round(3), "Table 2. Question category proportions.", floatfmt="{:.3f}")
    add_table_from_df_docx(doc, turn_summary.round(3), "Table 3. Turn-level proxy metrics.", floatfmt="{:.3f}")

    doc.add_heading("Figures", level=1)
    for caption, fn, width in [
        ("Figure 1. Turns per case by group and session.", "fig_turns_by_group_session.png", 5.8),
        ("Figure 2. Question category distribution.", "fig_question_categories.png", 6.5),
        ("Figure 3. Context grounding proxy (overlap).", "fig_context_overlap.png", 5.8),
    ]:
        p = doc.add_paragraph(caption)
        try:
            doc.add_picture(str(fig_dir / fn), width=Inches(width))
        except Exception:
            doc.add_paragraph(f"(Missing image: {fn})")
        doc.add_paragraph("")

    docx_path = out_dir / "MedSyn_Dialogue_Interaction_Report.docx"
    doc.save(docx_path)

    # PDF
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("MedSyn Interactive Dialogue Analysis (Session 2 & Session 4)", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Quantitative summaries", styles["Heading1"]))
    story.append(Paragraph("Table 1. Case-level summary.", styles["BodyText"]))
    story.append(df_to_rl_table(case_summary.round(3)))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Table 2. Question category proportions.", styles["BodyText"]))
    story.append(df_to_rl_table(q_tbl.round(3)))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Table 3. Turn-level proxy metrics.", styles["BodyText"]))
    story.append(df_to_rl_table(turn_summary.round(3)))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Figures", styles["Heading1"]))
    for caption, fn, width in [
        ("Figure 1. Turns per case by group and session.", "fig_turns_by_group_session.png", 6.0),
        ("Figure 2. Question category distribution.", "fig_question_categories.png", 6.5),
        ("Figure 3. Context grounding proxy (overlap).", "fig_context_overlap.png", 6.0),
    ]:
        story.append(Paragraph(caption, styles["BodyText"]))
        img_path = fig_dir / fn
        if img_path.exists():
            story.append(Image(str(img_path), width=width * inch, height=width * 0.6 * inch))
        else:
            story.append(Paragraph(f"(Missing image: {fn})", styles["BodyText"]))
        story.append(Spacer(1, 12))

    pdf_path = out_dir / "MedSyn_Dialogue_Interaction_Report.pdf"
    SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36).build(story)

    print(f"Wrote:\n- {docx_path}\n- {pdf_path}")


if __name__ == "__main__":
    main()
