import os, re, zipfile, subprocess
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

# Optional (for exploratory regressions)
try:
    import statsmodels.api as sm
    import statsmodels.formula.api as smf
    STATS = True
except Exception:
    STATS = False

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# -----------------------
# Input files (edit if needed)
# -----------------------
FILES = {
    "session1_baseline": "/home/gunel/medSyn/eval/session_outputs/auto_eval_session1_baseline.csv",
    "session2_interactive": "/home/gunel/medSyn/eval/session_outputs/auto_eval_session2_interactive.csv",
    "session3_baseline": "/home/gunel/medSyn/eval/session_outputs/auto_eval_session3_baseline.csv",
    "session4_interactive": "/home/gunel/medSyn/eval/session_outputs/auto_eval_session4_interactive.csv",
}

# Output paths
OUT_DIR = Path("/home/gunel/medSyn/eval/auto_eval")
PLOTS_DIR = OUT_DIR / "plots"
TABLES_DIR = OUT_DIR / "tables"
PLOTS_DIR.mkdir(parents=True, exist_ok=True)
TABLES_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = Path("MedSyn_Evaluation_Report_AnyMatch.docx")
PDF_PATH  = Path("MedSyn_Evaluation_Report_AnyMatch.pdf")
ZIP_PATH  = Path("MedSyn_Evaluation_AnyMatch_Package.zip")


# -----------------------
# Metric helpers
# -----------------------
def split_diag(s):
    """Split by ';' and trim around separators only. Does not modify cell content."""
    if s is None or (isinstance(s, float) and np.isnan(s)):
        return []
    s = str(s)
    if not s.strip():
        return []
    return [x.strip() for x in s.split(";") if x.strip()]

_TOKEN_RE = re.compile(r"[A-Za-z0-9]+", re.UNICODE)

def tokenize(s):
    if s is None or (isinstance(s, float) and np.isnan(s)):
        return []
    return _TOKEN_RE.findall(str(s).lower())

def any_match_accuracy(ans_items, gt_items):
    """Primary metric: 1 if any predicted diagnosis exactly matches any GT diagnosis."""
    if not ans_items or not gt_items:
        return 0
    return int(len(set(ans_items) & set(gt_items)) > 0)

def token_f1(ans_items, gt_items):
    """Secondary metric: list-level token overlap F1 (graded)."""
    ta = set(tokenize(" ".join(ans_items)))
    tg = set(tokenize(" ".join(gt_items)))
    if not ta and not tg:
        return 1.0
    if not ta or not tg:
        return 0.0
    inter = len(ta & tg)
    p = inter / len(ta)
    r = inter / len(tg)
    return 0.0 if (p + r) == 0 else 2 * p * r / (p + r)

def safe_float(x):
    try:
        if x is None:
            return np.nan
        if isinstance(x, str) and not x.strip():
            return np.nan
        return float(x)
    except Exception:
        return np.nan


# -----------------------
# Confidence intervals
# -----------------------
def cluster_boot_ci(df, value_col, cluster_col, n_boot=4000, seed=7):
    """Participant-cluster bootstrap CI for mean."""
    rng = np.random.default_rng(seed)
    clusters = df[cluster_col].dropna().unique()
    if len(clusters) == 0:
        return (np.nan, np.nan, np.nan)

    point = float(df[value_col].mean())
    boots = []
    for _ in range(n_boot):
        sampled = rng.choice(clusters, size=len(clusters), replace=True)
        boot_df = pd.concat([df[df[cluster_col] == c] for c in sampled], ignore_index=True)
        boots.append(float(boot_df[value_col].mean()))
    lo, hi = np.percentile(boots, [2.5, 97.5])
    return (point, float(lo), float(hi))

def boot_ci(values, n_boot=4000, seed=7):
    """Simple bootstrap CI for mean (used for participant-level case bootstrap)."""
    rng = np.random.default_rng(seed)
    vals = np.asarray(values, dtype=float)
    vals = vals[~np.isnan(vals)]
    if len(vals) == 0:
        return (np.nan, np.nan, np.nan)
    point = float(np.mean(vals))
    boots = []
    for _ in range(n_boot):
        sample = rng.choice(vals, size=len(vals), replace=True)
        boots.append(float(np.mean(sample)))
    lo, hi = np.percentile(boots, [2.5, 97.5])
    return (point, float(lo), float(hi))

def ci_str(t, is_time=False):
    m, lo, hi = t
    if np.isnan(m):
        return "NA"
    if is_time:
        return f"{m:.1f} [{lo:.1f}, {hi:.1f}]"
    return f"{m:.3f} [{lo:.3f}, {hi:.3f}]"


# -----------------------
# Plot helper
# -----------------------
def grouped_bar(df, metric, title, ylabel, out_path):
    plot = df.groupby(["condition", "expertise"])[metric].mean().reset_index()
    conds = ["baseline", "interactive"]
    exps = sorted(plot["expertise"].unique())
    x = np.arange(len(exps))
    w = 0.35

    fig, ax = plt.subplots(figsize=(7.2, 4.0))
    for i, cond in enumerate(conds):
        vals = []
        for exp in exps:
            s = plot[(plot.condition == cond) & (plot.expertise == exp)][metric]
            vals.append(float(s.iloc[0]) if len(s) else np.nan)
        ax.bar(x + (i - 0.5) * w, vals, w, label=cond)

    ax.set_xticks(x)
    ax.set_xticklabels(exps)
    ax.set_title(title)
    ax.set_ylabel(ylabel)
    ax.legend()
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)


# -----------------------
# DOCX->PDF conversion
# -----------------------
def convert_docx_to_pdf(docx_path, pdf_path):
    # Requires LibreOffice: soffice
    outdir = str(pdf_path.parent)
    profile = f"file:///tmp/lo_profile_{os.getpid()}_anymatch"
    subprocess.run(
        ["soffice", f"-env:UserInstallation={profile}", "--headless",
         "--convert-to", "pdf", "--outdir", outdir, str(docx_path)],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE
    )
    produced = Path(outdir) / (docx_path.stem + ".pdf")
    if produced.exists() and produced != pdf_path:
        produced.replace(pdf_path)


# -----------------------
# Main analysis
# -----------------------
def run():
    # Load & combine
    frames = []
    for key, path in FILES.items():
        df = pd.read_csv(path)
        df["session_key"] = key
        m = re.search(r"session(\d+)_", key)
        df["session_num"] = int(m.group(1)) if m else np.nan
        df["condition"] = "interactive" if "interactive" in key else "baseline"
        frames.append(df)
    wide = pd.concat(frames, ignore_index=True)

    # Identify columns
    gt_col = [c for c in wide.columns if c.strip().lower() == "discharge diagnosis"][0]
    diff_col = next((c for c in wide.columns if c.strip().lower() == "difficulty"), None)
    if diff_col is None:
        wide["Difficulty"] = "NA"
        diff_col = "Difficulty"

    # Participants inferred from *_answer columns
    answer_cols = [c for c in wide.columns if c.endswith("_answer")]
    participants = sorted({c[:-7] for c in answer_cols})
    participants = [p for p in participants if f"{p}_time" in wide.columns]

    # Expertise mapping (based on your IDs)
    expertise = {p: ("senior" if p.startswith("phy") else "resident") for p in participants}

    # Long-format
    recs = []
    for _, row in wide.iterrows():
        gt_items = split_diag(row[gt_col])
        for p in participants:
            ans_items = split_diag(row.get(f"{p}_answer", np.nan))
            recs.append({
                "session_key": row["session_key"],
                "session_num": int(row["session_num"]),
                "condition": row["condition"],
                "difficulty": row[diff_col],
                "participant": p,
                "expertise": expertise[p],
                "time_s": safe_float(row.get(f"{p}_time", np.nan)),
                "any_match_acc": any_match_accuracy(ans_items, gt_items),
                "token_f1": token_f1(ans_items, gt_items),
            })
    long_df = pd.DataFrame(recs)
    long_df.to_csv(TABLES_DIR / "long_format_metrics.csv", index=False)

    # Summaries (participant-cluster bootstrap)
    def summarize(df, group_cols):
        rows = []
        for keys, g in df.groupby(group_cols):
            if not isinstance(keys, tuple):
                keys = (keys,)
            r = dict(zip(group_cols, keys))
            r["n_participants"] = g["participant"].nunique()
            r["n_rows"] = len(g)
            r["any_match_acc_ci"] = ci_str(cluster_boot_ci(g, "any_match_acc", "participant"))
            r["token_f1_ci"] = ci_str(cluster_boot_ci(g, "token_f1", "participant"))
            r["mean_time_ci"] = ci_str(cluster_boot_ci(g.dropna(subset=["time_s"]), "time_s", "participant"), is_time=True)
            rows.append(r)
        return pd.DataFrame(rows)

    overall = summarize(long_df, ["condition"])
    by_exp  = summarize(long_df, ["condition", "expertise"])
    by_sess = summarize(long_df, ["session_num", "condition"])

    overall.to_csv(TABLES_DIR / "overall_by_condition.csv", index=False)
    by_exp.to_csv(TABLES_DIR / "by_condition_expertise.csv", index=False)
    by_sess.to_csv(TABLES_DIR / "by_session.csv", index=False)

    # Participant-focused (case bootstrap within each participant-condition)
    part_rows = []
    for p, g in long_df.groupby("participant"):
        for cond in ["baseline", "interactive"]:
            gg = g[g["condition"] == cond]
            part_rows.append({
                "participant": p,
                "expertise": expertise[p],
                "condition": cond,
                "any_match_acc_ci": ci_str(boot_ci(gg["any_match_acc"].values)),
                "token_f1_ci": ci_str(boot_ci(gg["token_f1"].values)),
                "mean_time_ci": ci_str(boot_ci(gg["time_s"].dropna().values), is_time=True),
                "n_rows": len(gg)
            })
    participant_summary = pd.DataFrame(part_rows)
    participant_summary.to_csv(TABLES_DIR / "participant_summary.csv", index=False)

    # Performance gap
    gap_rows = []
    for cond in ["baseline", "interactive"]:
        s = long_df[(long_df.condition == cond) & (long_df.expertise == "senior")]
        r = long_df[(long_df.condition == cond) & (long_df.expertise == "resident")]
        gap_rows.append({
            "condition": cond,
            "any_match_senior": float(s.any_match_acc.mean()),
            "any_match_resident": float(r.any_match_acc.mean()),
            "gap_senior_minus_resident": float(s.any_match_acc.mean() - r.any_match_acc.mean()),
            "time_senior": float(s.time_s.mean()),
            "time_resident": float(r.time_s.mean()),
            "gap_resident_minus_senior_time": float(r.time_s.mean() - s.time_s.mean()),
        })
    gap_df = pd.DataFrame(gap_rows)
    gap_df.to_csv(TABLES_DIR / "performance_gaps.csv", index=False)

    # Plots
    grouped_bar(long_df, "any_match_acc", "Any-match accuracy by condition and expertise", "Any-match accuracy",
                PLOTS_DIR / "anymatch_by_condition_expertise.png")
    grouped_bar(long_df, "token_f1", "Token-F1 by condition and expertise", "Token-F1",
                PLOTS_DIR / "tokenf1_by_condition_expertise.png")
    grouped_bar(long_df.dropna(subset=["time_s"]), "time_s", "Mean time by condition and expertise", "Time (s)",
                PLOTS_DIR / "time_by_condition_expertise.png")

    # Optional exploratory regressions (cluster-robust by participant)
    if STATS:
        reg_df = long_df.copy()
        reg_df["condition_bin"] = (reg_df.condition == "interactive").astype(int)
        reg_df["difficulty"] = reg_df["difficulty"].astype("category")
        reg_df["expertise"] = reg_df["expertise"].astype("category")
        reg_df["session_num"] = reg_df["session_num"].astype(int)

        # Logistic regression for any-match
        try:
            m = smf.glm(
                "any_match_acc ~ condition_bin * expertise + C(difficulty) + C(session_num)",
                data=reg_df, family=sm.families.Binomial()
            ).fit(cov_type="cluster", cov_kwds={"groups": reg_df["participant"]})

            rows = []
            for term in m.params.index:
                if term == "Intercept":
                    continue
                beta = float(m.params[term]); se = float(m.bse[term])
                lo = beta - 1.96 * se; hi = beta + 1.96 * se
                rows.append({
                    "term": term,
                    "odds_ratio": float(np.exp(beta)),
                    "ci_low": float(np.exp(lo)),
                    "ci_high": float(np.exp(hi)),
                    "p_value": float(m.pvalues[term]),
                })
            pd.DataFrame(rows).to_csv(TABLES_DIR / "reg_anymatch_logit_or.csv", index=False)
        except Exception:
            pass

    # Build DOCX report
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    r = title.add_run("MedSyn Evaluation Report (Any-match Accuracy)")
    r.bold = True
    r.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Primary metric: any-match accuracy (exact diagnosis string match after ';' split)")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    doc.add_heading("Primary metric definition", level=1)
    doc.add_paragraph(
        "Any-match accuracy = 1 if at least one diagnosis in the participant answer exactly matches "
        "at least one diagnosis in the ground truth (both split by ';', trimming whitespace around separators)."
    )

    doc.add_heading("Overall results", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Condition"
    t.rows[0].cells[1].text = "Any-match accuracy (95% CI)"
    t.rows[0].cells[2].text = "Token-F1 (95% CI)"
    t.rows[0].cells[3].text = "Mean time (s) (95% CI)"
    for _, row in overall.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(row["condition"])
        cells[1].text = str(row["any_match_acc_ci"])
        cells[2].text = str(row["token_f1_ci"])
        cells[3].text = str(row["mean_time_ci"])

    doc.add_heading("By expertise", level=1)
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Condition"
    t2.rows[0].cells[1].text = "Expertise"
    t2.rows[0].cells[2].text = "Any-match accuracy (95% CI)"
    t2.rows[0].cells[3].text = "Token-F1 (95% CI)"
    t2.rows[0].cells[4].text = "Mean time (s) (95% CI)"
    for _, row in by_exp.iterrows():
        cells = t2.add_row().cells
        cells[0].text = str(row["condition"])
        cells[1].text = str(row["expertise"])
        cells[2].text = str(row["any_match_acc_ci"])
        cells[3].text = str(row["token_f1_ci"])
        cells[4].text = str(row["mean_time_ci"])

    doc.add_heading("Figures", level=1)
    for fn, cap in [
        ("anymatch_by_condition_expertise.png", "Any-match accuracy by condition and expertise."),
        ("tokenf1_by_condition_expertise.png", "Token-F1 by condition and expertise."),
        ("time_by_condition_expertise.png", "Mean time by condition and expertise."),
    ]:
        doc.add_picture(str(PLOTS_DIR / fn), width=Inches(6.2))
        p = doc.add_paragraph(cap)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

    doc.add_heading("Limitations", level=1)
    doc.add_paragraph(
        "Any-match accuracy is lenient and can be inflated when ground truth contains multiple diagnoses. "
        "It should be interpreted as 'at least one correct diagnosis was identified'. Token-F1 is reported as a graded complement."
    )

    doc.save(DOCX_PATH)

    # Convert to PDF (requires LibreOffice)
    try:
        convert_docx_to_pdf(DOCX_PATH, PDF_PATH)
    except Exception:
        pass

    # Package everything
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as z:
        z.write(DOCX_PATH, arcname=DOCX_PATH.name)
        if PDF_PATH.exists():
            z.write(PDF_PATH, arcname=PDF_PATH.name)
        for f in OUT_DIR.rglob("*"):
            if f.is_file():
                z.write(f, arcname=str(f.relative_to(OUT_DIR.parent)))

    print("Done.")
    print("DOCX:", DOCX_PATH.resolve())
    print("PDF :", PDF_PATH.resolve(), "(if created)")
    print("ZIP :", ZIP_PATH.resolve())


if __name__ == "__main__":
    run()
