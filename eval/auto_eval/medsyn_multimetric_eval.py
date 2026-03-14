#!/usr/bin/env python3
"""
MedSyn evaluation (multi-metric) WITH difficulty stratification + session-paired analysis.

What it does
- Reads 4 session CSVs (S1..S4), where each clinician has columns:
    <id>_answer, <id>_time
  and there is:
    - Discharge diagnosis (ground truth; may include primary/secondary; we keep PRIMARY only)
    - Difficulty (case difficulty label; optional)

- Builds a long-format table: one row per (session, case, participant)
- Computes multiple metrics:
    * exact set match accuracy
    * any-match accuracy
    * precision / recall / F1 over diagnosis items
    * Jaccard over diagnosis items
    * ROUGE-1 F1, ROUGE-L F1 (token-based)
    * TF-IDF cosine
    * BERTScore F1 (optional; if bert-score installed and model available)

- Summaries with 95% bootstrap CIs at the PARTICIPANT level:
    * session overall
    * session by expertise (senior vs resident)
    * session by difficulty
    * session by difficulty × expertise

- Paired improvements:
    * S2−S1 and S4−S3 (participant-level paired)
    * overall, by expertise
    * by difficulty, by difficulty × expertise

- Optional report outputs:
    * DOCX (python-docx)
    * PDF companion (reportlab)
    * LaTeX .tex (tables+figures)
    * ZIP package of everything

Assumptions
- Time columns are in seconds (we keep seconds; also show minutes in some report sections).
- Diagnosis “items” are split on ';' and newlines.
- Ground truth may include primary/secondary; we attempt to extract PRIMARY only.

Usage example
python medsyn_multimetric_eval_with_difficulty.py \
  --s1 auto_eval_session1_baseline.csv \
  --s2 auto_eval_session2_interactive.csv \
  --s3 auto_eval_session3_baseline.csv \
  --s4 auto_eval_session4_interactive.csv \
  --out medsyn_eval_out \
  --make-docx --make-pdf --make-tex --make-zip
"""

from __future__ import annotations

import argparse
import re
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd

# plots are optional but enabled by default
import matplotlib.pyplot as plt

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# -----------------------------
# Configuration defaults
# -----------------------------
DEFAULT_SESSION_COND = {1: "baseline", 2: "interactive", 3: "baseline", 4: "interactive"}
DEFAULT_PAIRINGS = [(1, 2), (3, 4)]

_TOKEN_RE = re.compile(r"[A-Za-z0-9]+", re.UNICODE)


# -----------------------------
# Parsing ground truth / answers
# -----------------------------
def _norm_str(x) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return ""
    return str(x).strip()


def extract_primary_text(gt_raw: str) -> str:
    """
    Attempt to extract PRIMARY diagnoses from a discharge diagnosis cell.

    Handles common patterns:
    - "Primary diagnosis: .... \n Secondary diagnosis: ...."
    - "Primary diagnosis ...." (no secondary)
    - If no labels: cut at a line that looks like "Secondary diagnosis"
    - Keep multiple primary diagnoses if they are on separate lines
    """
    s = _norm_str(gt_raw)
    if not s:
        return ""

    # Normalize line endings
    s = s.replace("\r\n", "\n").replace("\r", "\n")

    # If there is an explicit "Primary diagnosis" label, keep text after it
    m = re.search(r"primary\s*diagnosis\s*:?\s*(.*)$", s, flags=re.IGNORECASE | re.DOTALL)
    if m:
        s = m.group(1).strip()

    # Cut at "Secondary diagnosis" label if present
    s = re.split(r"\n\s*secondary\s*diagnosis\s*:?", s, flags=re.IGNORECASE)[0].strip()

    # Some notes put secondary after a blank line + header; also cut at "Secondary diagnoses"
    s = re.split(r"\n\s*secondary\s*diagnoses\s*:?", s, flags=re.IGNORECASE)[0].strip()

    return s


def split_diagnosis_items(text: str) -> List[str]:
    """
    Split diagnoses into 'items' using:
      - semicolon ';'
      - newline '\n'
    """
    s = _norm_str(text)
    if not s:
        return []
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    parts = re.split(r"[;\n]+", s)
    out = [p.strip() for p in parts if p.strip()]
    return out


def tokens(text: str) -> List[str]:
    s = _norm_str(text).lower()
    if not s:
        return []
    return _TOKEN_RE.findall(s)

def _token_set(text: str) -> set:
    """Token set for containment matching (case-insensitive alphanumerics)."""
    return set(tokens(text))


def _count_containment_matches(pred_items: List[str], gt_items: List[str]) -> int:
    """Count 1-to-1 containment matches between predicted and ground-truth diagnosis items.

    A predicted diagnosis item is considered correct if ALL its tokens are a subset of the
    tokens of some ground-truth diagnosis item. Each GT item can be matched at most once
    (greedy matching), preventing double-counting.
    """
    if not pred_items or not gt_items:
        return 0

    pred_tok = [(i, _token_set(pred_items[i])) for i in range(len(pred_items))]
    gt_tok = [(j, _token_set(gt_items[j])) for j in range(len(gt_items))]

    # Drop empty-token items (rare but possible)
    pred_tok = [(i, s) for i, s in pred_tok if len(s) > 0]
    gt_tok = [(j, s) for j, s in gt_tok if len(s) > 0]
    if not pred_tok or not gt_tok:
        return 0

    # Greedy: match more specific predictions first (more tokens),
    # and pick the smallest GT that contains them to avoid wasting broad GT items.
    unmatched_gt = {j for j, _ in gt_tok}
    gt_map = {j: s for j, s in gt_tok}

    matches = 0
    pred_tok_sorted = sorted(pred_tok, key=lambda x: (-len(x[1]), str(pred_items[x[0]]).lower()))
    for i, pset in pred_tok_sorted:
        candidates = []
        for j in list(unmatched_gt):
            gset = gt_map[j]
            if pset.issubset(gset):
                candidates.append((len(gset), j))
        if candidates:
            _, best_j = min(candidates, key=lambda t: t[0])
            unmatched_gt.remove(best_j)
            matches += 1

    return matches



def safe_float(x) -> float:
    try:
        if x is None:
            return np.nan
        if isinstance(x, str) and not x.strip():
            return np.nan
        return float(x)
    except Exception:
        return np.nan


# -----------------------------
# Metrics over diagnosis items
# -----------------------------
def exact_set_acc(pred_items: List[str], gt_items: List[str]) -> int:
    if not pred_items and not gt_items:
        return 1
    return int(set(pred_items) == set(gt_items))


def any_match_acc(pred_items: List[str], gt_items: List[str]) -> int:
    """Lenient accuracy: 1 if any predicted item is contained in any GT item.

    Containment rule: tokens(pred_item) ⊆ tokens(gt_item).
    """
    return int(_count_containment_matches(pred_items, gt_items) > 0)

def precision_recall_f1(pred_items: List[str], gt_items: List[str]) -> Tuple[float, float, float]:
    """Precision/recall/F1 over diagnosis items using containment matching.

    True positives are counted via 1-to-1 greedy containment matches:
      pred item is correct if tokens(pred) ⊆ tokens(gt item).
    """
    # Degenerate 'both empty' case as perfect.
    if (not pred_items) and (not gt_items):
        return 1.0, 1.0, 1.0

    tp = _count_containment_matches(pred_items, gt_items)

    # Use raw non-empty item counts for denominators.
    n_pred = len([p for p in pred_items if _norm_str(p)])
    n_gt = len([g for g in gt_items if _norm_str(g)])

    prec = tp / n_pred if n_pred else 0.0
    rec = tp / n_gt if n_gt else 0.0
    f1 = 0.0 if (prec + rec) == 0 else (2 * prec * rec / (prec + rec))
    return prec, rec, f1

def jaccard_items(pred_items: List[str], gt_items: List[str]) -> float:
    """Jaccard over diagnosis items under containment matching.

    We compute a count-based Jaccard: TP / (|pred| + |gt| − TP),
    where TP is the 1-to-1 greedy containment match count.
    """
    if (not pred_items) and (not gt_items):
        return 1.0
    tp = _count_containment_matches(pred_items, gt_items)
    n_pred = len([p for p in pred_items if _norm_str(p)])
    n_gt = len([g for g in gt_items if _norm_str(g)])
    denom = (n_pred + n_gt - tp)
    return (tp / denom) if denom else 0.0


# -----------------------------
# Soft text similarity metrics
# -----------------------------
def rouge1_f1(pred_text: str, gt_text: str) -> float:
    pt = tokens(pred_text)
    gt = tokens(gt_text)
    if not pt and not gt:
        return 1.0
    if not pt or not gt:
        return 0.0
    from collections import Counter

    pc, gc = Counter(pt), Counter(gt)
    overlap = sum((pc & gc).values())
    prec = overlap / sum(pc.values())
    rec = overlap / sum(gc.values())
    return 0.0 if (prec + rec) == 0 else 2 * prec * rec / (prec + rec)


def lcs_len(a: List[str], b: List[str]) -> int:
    n, m = len(a), len(b)
    if n == 0 or m == 0:
        return 0
    dp = [0] * (m + 1)
    for i in range(1, n + 1):
        prev = 0
        for j in range(1, m + 1):
            tmp = dp[j]
            if a[i - 1] == b[j - 1]:
                dp[j] = prev + 1
            else:
                dp[j] = max(dp[j], dp[j - 1])
            prev = tmp
    return dp[m]


def rougeL_f1(pred_text: str, gt_text: str) -> float:
    pt = tokens(pred_text)
    gt = tokens(gt_text)
    if not pt and not gt:
        return 1.0
    if not pt or not gt:
        return 0.0
    lcs = lcs_len(pt, gt)
    prec = lcs / len(pt)
    rec = lcs / len(gt)
    return 0.0 if (prec + rec) == 0 else 2 * prec * rec / (prec + rec)


def tfidf_cosine(pred_text: str, gt_text: str) -> float:
    vec = TfidfVectorizer(token_pattern=r"(?u)\b\w+\b", lowercase=True)
    X = vec.fit_transform([pred_text or "", gt_text or ""])
    return float(cosine_similarity(X[0], X[1])[0, 0])


def try_bertscore(pred_texts: List[str], gt_texts: List[str]) -> Tuple[np.ndarray, bool]:
    """
    Optional. If unavailable, returns NaNs and bs_ok=False.
    """
    try:
        from bert_score import score as bert_score  # type: ignore

        _, _, F1 = bert_score(
            pred_texts,
            gt_texts,
            lang="en",
            verbose=False,
            rescale_with_baseline=False,
        )
        return np.array(F1.cpu()), True
    except Exception:
        return np.array([np.nan] * len(pred_texts)), False


# -----------------------------
# Bootstrap summaries (participant-level)
# -----------------------------
def bootstrap_mean_ci(values: np.ndarray, n_boot: int, seed: int) -> Tuple[float, float, float]:
    v = np.asarray(values, dtype=float)
    v = v[~np.isnan(v)]
    if len(v) == 0:
        return (np.nan, np.nan, np.nan)
    point = float(np.mean(v))
    if len(v) == 1:
        return (point, point, point)
    rng = np.random.default_rng(seed)
    boots = rng.choice(v, size=(n_boot, len(v)), replace=True).mean(axis=1)
    lo, hi = np.percentile(boots, [2.5, 97.5])
    return (point, float(lo), float(hi))


def fmt_ci(t: Tuple[float, float, float], is_time: bool = False) -> str:
    m, lo, hi = t
    if np.isnan(m):
        return "NA"
    return f"{m:.1f} [{lo:.1f}, {hi:.1f}]" if is_time else f"{m:.3f} [{lo:.3f}, {hi:.3f}]"


def infer_difficulty_order(values: List[str]) -> List[str]:
    vals = sorted(set([v for v in values if _norm_str(v)]))
    norm = [v.strip().lower() for v in vals]

    # common patterns
    for pref in [["easy", "medium", "hard"], ["low", "medium", "high"], ["1", "2", "3", "4", "5"]]:
        if all(p in norm for p in pref):
            order = [vals[norm.index(p)] for p in pref]
            # append leftovers (if any)
            for v in vals:
                if v not in order:
                    order.append(v)
            return order
    return vals


def sort_by_difficulty(df: pd.DataFrame, difficulty_order: List[str], col: str = "difficulty") -> pd.DataFrame:
    dmap = {d: i for i, d in enumerate(difficulty_order)}
    out = df.copy()
    out[col] = out[col].astype(str)
    out["_difficulty_order"] = out[col].map(lambda x: dmap.get(x, 999))
    out = out.sort_values(["_difficulty_order"]).drop(columns=["_difficulty_order"])
    return out


# -----------------------------
# DOCX / PDF / TEX helpers (optional)
# -----------------------------
def _try_import_docx():
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.oxml.ns import qn  # type: ignore

        return Document, Inches, Pt, WD_ALIGN_PARAGRAPH, qn
    except Exception:
        return None


def _try_import_reportlab():
    try:
        from reportlab.lib.pagesizes import letter  # type: ignore
        from reportlab.pdfgen import canvas  # type: ignore
        from reportlab.lib.units import inch  # type: ignore
        from reportlab.lib.utils import ImageReader  # type: ignore

        return letter, canvas, inch, ImageReader
    except Exception:
        return None


def latex_escape(s: str) -> str:
    s = str(s)
    s = s.replace("\\", "\\textbackslash ")
    for ch in ["&", "%", "$", "#", "_", "{", "}", "~", "^"]:
        s = s.replace(ch, "\\" + ch)
    return s


def df_to_latex(df: pd.DataFrame, caption: str, label: str) -> str:
    cols = list(df.columns)
    colspec = "l" * len(cols)
    lines = []
    lines.append("\\begin{table}[t]")
    lines.append("\\centering")
    lines.append(f"\\caption{{{latex_escape(caption)}}}")
    lines.append(f"\\label{{{label}}}")
    lines.append(f"\\begin{{tabular}}{{{colspec}}}")
    lines.append("\\toprule")
    lines.append(" & ".join(latex_escape(c) for c in cols) + " \\\\")
    lines.append("\\midrule")
    for _, r in df.iterrows():
        lines.append(" & ".join(latex_escape(r[c]) for c in cols) + " \\\\")
    lines.append("\\bottomrule")
    lines.append("\\end{tabular}")
    lines.append("\\end{table}")
    return "\n".join(lines)


# -----------------------------
# Core pipeline
# -----------------------------
@dataclass
class Outputs:
    out_dir: Path
    tables_dir: Path
    plots_dir: Path
    long_path: Path


def build_long_df(
    session_dfs: Dict[int, pd.DataFrame],
    session_cond: Dict[int, str],
    gt_col: str,
    diff_col: str,
    participants: List[str],
    expertise: Dict[str, str],
) -> pd.DataFrame:
    rows = []
    for s, df in session_dfs.items():
        for case_idx, row in df.iterrows():
            gt_primary = extract_primary_text(row[gt_col])
            gt_items = split_diagnosis_items(gt_primary)
            gt_text = " ; ".join(gt_items)

            difficulty = row.get(diff_col, "NA")

            for p in participants:
                pred_raw = row.get(f"{p}_answer", np.nan)
                pred_items = split_diagnosis_items(pred_raw)
                pred_text = " ; ".join(pred_items)

                prec, rec, f1 = precision_recall_f1(pred_items, gt_items)

                rows.append(
                    {
                        "session": s,
                        "condition": session_cond[s],
                        "case_in_session": int(case_idx + 1),
                        "difficulty": difficulty,
                        "participant": p,
                        "expertise": expertise.get(p, "unknown"),
                        "time_s": safe_float(row.get(f"{p}_time", np.nan)),
                        "time_min": safe_float(row.get(f"{p}_time", np.nan)) / 60.0 if safe_float(row.get(f"{p}_time", np.nan)) == safe_float(row.get(f"{p}_time", np.nan)) else np.nan,
                        "n_pred_diags": int(len(pred_items)),
                        "gt_raw": row[gt_col],
                        "gt_primary": gt_primary,
                        "pred_raw": pred_raw,
                        "pred_text": pred_text,
                        "gt_text": gt_text,
                        "exact_set_acc": exact_set_acc(pred_items, gt_items),
                        "any_match_acc": any_match_acc(pred_items, gt_items),
                        "precision_items": prec,
                        "recall_items": rec,
                        "f1_items": f1,
                        "jaccard_items": jaccard_items(pred_items, gt_items),
                        "rouge1_f1": rouge1_f1(pred_text, gt_text),
                        "rougeL_f1": rougeL_f1(pred_text, gt_text),
                        "tfidf_cosine": tfidf_cosine(pred_text, gt_text),
                    }
                )

    long_df = pd.DataFrame(rows)
    return long_df


def summarize_participant_bootstrap(
    df: pd.DataFrame,
    group_cols: List[str],
    metrics: List[str],
    n_boot: int,
    seed: int,
) -> pd.DataFrame:
    out = []
    for keys, g in df.groupby(group_cols):
        if not isinstance(keys, tuple):
            keys = (keys,)
        row = dict(zip(group_cols, keys))

        row["n_participants"] = int(g["participant"].nunique())
        row["n_rows"] = int(len(g))

        part = g.groupby("participant")
        for m in metrics:
            row[m + "_ci"] = fmt_ci(bootstrap_mean_ci(part[m].mean().values, n_boot=n_boot, seed=seed))

        row["mean_time_s_ci"] = fmt_ci(bootstrap_mean_ci(part["time_s"].mean().values, n_boot=n_boot, seed=seed), is_time=True)
        row["mean_time_min_ci"] = fmt_ci(bootstrap_mean_ci(part["time_min"].mean().values, n_boot=n_boot, seed=seed), is_time=True)
        row["mean_n_pred_ci"] = fmt_ci(bootstrap_mean_ci(part["n_pred_diags"].mean().values, n_boot=n_boot, seed=seed))

        out.append(row)
    return pd.DataFrame(out)


def participant_means(df: pd.DataFrame, session: int, metrics: List[str]) -> pd.DataFrame:
    sub = df[df.session == session]
    agg = {m: (m, "mean") for m in metrics}
    return sub.groupby("participant").agg(
        **agg,
        mean_time=("time_s", "mean"),
        mean_time_min=("time_min", "mean"),
        mean_n_pred=("n_pred_diags", "mean"),
    )


def participant_means_by_diff(df: pd.DataFrame, session: int, metrics: List[str]) -> pd.DataFrame:
    sub = df[df.session == session]
    agg = {m: (m, "mean") for m in metrics}
    return (
        sub.groupby(["participant", "difficulty"])
        .agg(
            **agg,
            mean_time=("time_s", "mean"),
            mean_time_min=("time_min", "mean"),
            mean_n_pred=("n_pred_diags", "mean"),
        )
        .reset_index()
    )


def summarize_improvements_bootstrap(
    df: pd.DataFrame,
    group_cols: List[str],
    metrics: List[str],
    n_boot: int,
    seed: int,
) -> pd.DataFrame:
    out = []
    for keys, g in df.groupby(group_cols):
        if not isinstance(keys, tuple):
            keys = (keys,)
        row = dict(zip(group_cols, keys))
        row["n_participants"] = int(len(g))

        for m in metrics:
            row[m + "_diff_ci"] = fmt_ci(bootstrap_mean_ci(g[m].values, n_boot=n_boot, seed=seed))
        row["mean_time_diff_ci"] = fmt_ci(bootstrap_mean_ci(g["mean_time"].values, n_boot=n_boot, seed=seed), is_time=True)
        row["mean_time_min_diff_ci"] = fmt_ci(bootstrap_mean_ci(g["mean_time_min"].values, n_boot=n_boot, seed=seed), is_time=True)
        row["mean_n_pred_diff_ci"] = fmt_ci(bootstrap_mean_ci(g["mean_n_pred"].values, n_boot=n_boot, seed=seed))

        out.append(row)
    return pd.DataFrame(out)


def within_session_consensus_item_f1(long_df: pd.DataFrame) -> pd.DataFrame:
    """
    Consensus: for each session, compute mean pairwise F1 among participants per case,
    then average across cases.
    """
    def f1_from_sets(a: set, b: set) -> float:
        if not a and not b:
            return 1.0
        if not a or not b:
            return 0.0
        inter = len(a & b)
        return 2 * inter / (len(a) + len(b)) if (len(a) + len(b)) else 0.0

    rows = []
    for s in sorted(long_df.session.unique()):
        sub = long_df[long_df.session == s]
        case_vals = []
        for case in sorted(sub.case_in_session.unique()):
            cc = sub[sub.case_in_session == case]
            sets = {r["participant"]: set(split_diagnosis_items(r["pred_raw"])) for _, r in cc.iterrows()}
            parts = list(sets.keys())
            vals = []
            for i in range(len(parts)):
                for j in range(i + 1, len(parts)):
                    vals.append(f1_from_sets(sets[parts[i]], sets[parts[j]]))
            case_vals.append(float(np.mean(vals)) if vals else np.nan)
        rows.append(
            {
                "session": int(s),
                "pairwise_consensus_itemF1": float(np.nanmean(case_vals)),
            }
        )
    return pd.DataFrame(rows)


def make_plots(long_df: pd.DataFrame, plots_dir: Path, difficulty_order: List[str]) -> None:
    plots_dir.mkdir(parents=True, exist_ok=True)

    def line_by_session(metric: str, title: str, ylabel: str, out_png: Path) -> None:
        grp = long_df.groupby(["session", "expertise"])[metric].mean().reset_index().sort_values("session")
        fig, ax = plt.subplots(figsize=(8.0, 4.4))
        for exp in sorted(grp.expertise.unique()):
            sub = grp[grp.expertise == exp]
            ax.plot(sub.session, sub[metric], marker="o", label=exp)
        ax.set_xticks(sorted(grp.session.unique()))
        ax.set_title(title)
        ax.set_xlabel("Session")
        ax.set_ylabel(ylabel)
        ax.legend()
        fig.tight_layout()
        fig.savefig(out_png, dpi=200)
        plt.close(fig)

    line_by_session("any_match_acc", "Any-match accuracy across sessions", "Any-match accuracy", plots_dir / "anymatch_across_sessions.png")
    line_by_session("f1_items", "Diagnosis-item F1 across sessions", "Item F1", plots_dir / "itemf1_across_sessions.png")
    line_by_session("time_s", "Mean time across sessions", "Time (s)", plots_dir / "time_across_sessions.png")

    # Paired improvement by difficulty (mean delta item F1) if we can compute it later in tables
    # (We create it from long_df directly here as a visual helper)
    # This is not the CI version; it's just the mean delta curve.
    # It requires that sessions 1/2 and 3/4 share the same difficulty labels.
    try:
        # compute participant means by difficulty, then paired delta, then average
        metrics = ["f1_items"]
        pair_dfs = []
        for a, b in DEFAULT_PAIRINGS:
            A = participant_means_by_diff(long_df, a, metrics)
            B = participant_means_by_diff(long_df, b, metrics)
            merged = A.merge(B, on=["participant", "difficulty"], suffixes=("_A", "_B"))
            merged["pair"] = f"S{a}->S{b}"
            merged["delta_f1_items"] = merged["f1_items_B"] - merged["f1_items_A"]
            pair_dfs.append(merged[["pair", "difficulty", "delta_f1_items"]])

        dd = pd.concat(pair_dfs, ignore_index=True)
        dd["difficulty"] = dd["difficulty"].astype(str)
        # order difficulty
        dmap = {d: i for i, d in enumerate(difficulty_order)}
        dd["_do"] = dd["difficulty"].map(lambda x: dmap.get(x, 999))
        dd = dd.sort_values(["pair", "_do"]).drop(columns=["_do"])
        diffs_sorted = [d for d in difficulty_order if d in dd["difficulty"].unique()]

        fig, ax = plt.subplots(figsize=(8.2, 4.6))
        for pair in sorted(dd["pair"].unique()):
            sub = dd[dd["pair"] == pair].groupby("difficulty")["delta_f1_items"].mean().reindex(diffs_sorted)
            ax.plot(range(len(diffs_sorted)), sub.values, marker="o", label=pair)
        ax.axhline(0, linestyle="--")
        ax.set_xticks(range(len(diffs_sorted)))
        ax.set_xticklabels(diffs_sorted, rotation=25, ha="right")
        ax.set_title("Paired improvement in item-F1 by difficulty (mean Δ)")
        ax.set_xlabel("Difficulty")
        ax.set_ylabel("Δ item-F1 (interactive − baseline)")
        ax.legend()
        fig.tight_layout()
        fig.savefig(plots_dir / "paired_itemf1_by_difficulty_mean.png", dpi=200)
        plt.close(fig)
    except Exception:
        pass


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--s1", required=True, help="Session 1 CSV (baseline)")
    ap.add_argument("--s2", required=True, help="Session 2 CSV (interactive)")
    ap.add_argument("--s3", required=True, help="Session 3 CSV (baseline)")
    ap.add_argument("--s4", required=True, help="Session 4 CSV (interactive)")
    ap.add_argument("--out", default="/home/gunel/medSyn/eval/auto_eval", help="Output folder")
    ap.add_argument("--seed", type=int, default=7)
    ap.add_argument("--n-boot", type=int, default=4000)

    ap.add_argument("--make-docx", action="store_true")
    ap.add_argument("--make-pdf", action="store_true")
    ap.add_argument("--make-tex", action="store_true")
    ap.add_argument("--make-zip", action="store_true")

    args = ap.parse_args()

    seed = args.seed
    n_boot = args.n_boot

    files = {1: args.s1, 2: args.s2, 3: args.s3, 4: args.s4}
    session_cond = DEFAULT_SESSION_COND

    out_dir = Path(args.out)
    tables_dir = out_dir / "tables"
    plots_dir = out_dir / "plots"
    out_dir.mkdir(parents=True, exist_ok=True)
    tables_dir.mkdir(exist_ok=True)
    plots_dir.mkdir(exist_ok=True)

    # read
    session_dfs = {s: pd.read_csv(p) for s, p in files.items()}
    wide_any = session_dfs[1]

    # columns
    gt_candidates = [c for c in wide_any.columns if c.strip().lower() == "discharge diagnosis"]
    if not gt_candidates:
        raise ValueError("Could not find 'Discharge diagnosis' column (case-insensitive match).")
    gt_col = gt_candidates[0]

    diff_col = next((c for c in wide_any.columns if c.strip().lower() == "difficulty"), None)
    if diff_col is None:
        diff_col = "Difficulty"
        for s in session_dfs:
            session_dfs[s][diff_col] = "NA"

    answer_cols = [c for c in wide_any.columns if c.endswith("_answer")]
    participants = sorted({c[:-7] for c in answer_cols})
    participants = [p for p in participants if f"{p}_time" in wide_any.columns]

    # expertise heuristic (your convention)
    expertise = {p: ("senior" if p.startswith("phy") else "resident" if p.startswith("res") else "unknown") for p in participants}

    # infer difficulty order from all sessions
    diff_values = []
    for s in session_dfs:
        diff_values.extend(list(session_dfs[s][diff_col].dropna().astype(str).unique()))
    difficulty_order = infer_difficulty_order(diff_values)

    # build long df
    long_df = build_long_df(session_dfs, session_cond, gt_col, diff_col, participants, expertise)

    # optional BERTScore
    bs_f1, bs_ok = try_bertscore(long_df["pred_text"].tolist(), long_df["gt_text"].tolist())
    long_df["bertscore_f1"] = bs_f1

    # metrics list
    METRICS = [
        "exact_set_acc",
        "any_match_acc",
        "precision_items",
        "recall_items",
        "f1_items",
        "jaccard_items",
        "rouge1_f1",
        "rougeL_f1",
        "tfidf_cosine",
        "bertscore_f1",
    ]

    # save long
    long_path = tables_dir / "long_format_multimetrics.csv"
    long_df.to_csv(long_path, index=False)

    # summaries
    sess_overall = summarize_participant_bootstrap(long_df, ["session", "condition"], METRICS, n_boot, seed).sort_values("session")
    sess_overall.to_csv(tables_dir / "session_overall_summary.csv", index=False)

    sess_by_exp = summarize_participant_bootstrap(long_df, ["session", "condition", "expertise"], METRICS, n_boot, seed).sort_values(["session", "expertise"])
    sess_by_exp.to_csv(tables_dir / "session_by_expertise_summary.csv", index=False)

    sess_by_diff = summarize_participant_bootstrap(long_df, ["session", "condition", "difficulty"], METRICS, n_boot, seed)
    sess_by_diff = sess_by_diff.sort_values(["session"]).reset_index(drop=True)
    # sort difficulties within each session for readability
    out_parts = []
    for s in sorted(sess_by_diff["session"].unique()):
        part = sess_by_diff[sess_by_diff["session"] == s]
        out_parts.append(sort_by_difficulty(part, difficulty_order, col="difficulty"))
    sess_by_diff = pd.concat(out_parts, ignore_index=True)
    sess_by_diff.to_csv(tables_dir / "session_by_difficulty_summary.csv", index=False)

    sess_by_diff_exp = summarize_participant_bootstrap(long_df, ["session", "condition", "difficulty", "expertise"], METRICS, n_boot, seed)
    out_parts = []
    for s in sorted(sess_by_diff_exp["session"].unique()):
        part = sess_by_diff_exp[sess_by_diff_exp["session"] == s]
        part2 = sort_by_difficulty(part, difficulty_order, col="difficulty")
        out_parts.append(part2.sort_values(["difficulty", "expertise"]))
    sess_by_diff_exp = pd.concat(out_parts, ignore_index=True)
    sess_by_diff_exp.to_csv(tables_dir / "session_by_difficulty_expertise_summary.csv", index=False)

    # paired improvements overall/by expertise
    pair_parts = []
    for a, b in DEFAULT_PAIRINGS:
        A = participant_means(long_df, a, METRICS)
        B = participant_means(long_df, b, METRICS)
        common = A.index.intersection(B.index)
        D = B.loc[common] - A.loc[common]
        D["pair"] = f"S{a}->S{b}"
        D["participant"] = D.index
        D["expertise"] = [expertise.get(p, "unknown") for p in D.index]
        pair_parts.append(D.reset_index(drop=True))

    pair_improv = pd.concat(pair_parts, ignore_index=True)
    pair_improv.to_csv(tables_dir / "paired_improvements_by_participant.csv", index=False)

    improv_overall = summarize_improvements_bootstrap(pair_improv, ["pair"], METRICS, n_boot, seed).sort_values("pair")
    improv_overall.to_csv(tables_dir / "paired_improvements_overall.csv", index=False)

    improv_by_exp = summarize_improvements_bootstrap(pair_improv, ["pair", "expertise"], METRICS, n_boot, seed).sort_values(["pair", "expertise"])
    improv_by_exp.to_csv(tables_dir / "paired_improvements_by_expertise.csv", index=False)

    # paired improvements by difficulty (+ expertise)
    pair_diff_parts = []
    for a, b in DEFAULT_PAIRINGS:
        A = participant_means_by_diff(long_df, a, METRICS)
        B = participant_means_by_diff(long_df, b, METRICS)
        merged = A.merge(B, on=["participant", "difficulty"], suffixes=("_A", "_B"))

        out = {
            "pair": f"S{a}->S{b}",
            "participant": merged["participant"],
            "difficulty": merged["difficulty"].astype(str),
            "expertise": merged["participant"].map(expertise),
        }
        for m in METRICS:
            out[m] = merged[f"{m}_B"] - merged[f"{m}_A"]
        out["mean_time"] = merged["mean_time_B"] - merged["mean_time_A"]
        out["mean_time_min"] = merged["mean_time_min_B"] - merged["mean_time_min_A"]
        out["mean_n_pred"] = merged["mean_n_pred_B"] - merged["mean_n_pred_A"]

        pair_diff_parts.append(pd.DataFrame(out))

    pair_improv_by_diff = pd.concat(pair_diff_parts, ignore_index=True)
    pair_improv_by_diff.to_csv(tables_dir / "paired_improvements_by_difficulty_participant.csv", index=False)

    improv_by_diff = summarize_improvements_bootstrap(pair_improv_by_diff, ["pair", "difficulty"], METRICS, n_boot, seed)
    out_parts = []
    for pair in sorted(improv_by_diff["pair"].unique()):
        part = improv_by_diff[improv_by_diff["pair"] == pair]
        out_parts.append(sort_by_difficulty(part, difficulty_order, col="difficulty"))
    improv_by_diff = pd.concat(out_parts, ignore_index=True)
    improv_by_diff.to_csv(tables_dir / "paired_improvements_by_difficulty.csv", index=False)

    improv_by_diff_exp = summarize_improvements_bootstrap(pair_improv_by_diff, ["pair", "difficulty", "expertise"], METRICS, n_boot, seed)
    out_parts = []
    for pair in sorted(improv_by_diff_exp["pair"].unique()):
        part = improv_by_diff_exp[improv_by_diff_exp["pair"] == pair]
        part = sort_by_difficulty(part, difficulty_order, col="difficulty").sort_values(["difficulty", "expertise"])
        out_parts.append(part)
    improv_by_diff_exp = pd.concat(out_parts, ignore_index=True)
    improv_by_diff_exp.to_csv(tables_dir / "paired_improvements_by_difficulty_expertise.csv", index=False)

    # consensus
    cons_df = within_session_consensus_item_f1(long_df)
    cons_df = cons_df.merge(pd.DataFrame({"session": [1, 2, 3, 4], "condition": [session_cond[1], session_cond[2], session_cond[3], session_cond[4]]}), on="session", how="left")
    cons_df.to_csv(tables_dir / "consensus_by_session.csv", index=False)

    # plots
    make_plots(long_df, plots_dir, difficulty_order)

    # optional report generation
    report_docx = out_dir / "MedSyn_Evaluation_Report_MultiMetric_WithDifficulty.docx"
    report_pdf = out_dir / "MedSyn_Evaluation_Report_MultiMetric_WithDifficulty.pdf"
    report_tex = out_dir / "MedSyn_Evaluation_Report_MultiMetric_WithDifficulty.tex"
    report_zip = out_dir / "MedSyn_MultiMetric_Evaluation_WithDifficulty_Package.zip"

    if args.make_docx:
        docx_mod = _try_import_docx()
        if docx_mod is None:
            print("DOCX requested but python-docx not available. Install: pip install python-docx")
        else:
            Document, Inches, Pt, WD_ALIGN_PARAGRAPH, qn = docx_mod

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
            style.font.size = Pt(11)

            p = doc.add_paragraph()
            rr = p.add_run("MedSyn Evaluation Report (Multi-metric + Difficulty Stratification, Session-paired)")
            rr.bold = True
            rr.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = doc.add_paragraph(
                f"Bootstrapping: participant-level means, 95% CI, seed={seed}, bootstraps={n_boot}. "
                f"Time is treated as seconds (also reported in minutes). "
                f"BERTScore computed: {bs_ok}."
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True

            doc.add_heading("Approach", level=1)
            doc.add_paragraph(
                "We evaluate each session separately, then quantify LLM assistance using paired differences "
                "(S2−S1 and S4−S3). We stratify results by clinician expertise and case difficulty."
            )
            doc.add_paragraph(
                "Ground truth: we extract PRIMARY diagnoses from the discharge diagnosis field (if primary/secondary are present), "
                "then split diagnosis items on ';' and newlines."
            )

            def add_table(df: pd.DataFrame, title: str, max_rows: int = 60):
                doc.add_paragraph(title).runs[0].bold = True
                show = df.copy()
                if len(show) > max_rows:
                    show = show.head(max_rows)
                    doc.add_paragraph(f"(Showing first {max_rows} rows; full table saved in CSV exports.)").runs[0].italic = True
                t = doc.add_table(rows=1, cols=len(show.columns))
                t.style = "Table Grid"
                for j, c in enumerate(show.columns):
                    t.rows[0].cells[j].text = str(c)
                for _, r in show.iterrows():
                    cells = t.add_row().cells
                    for j, c in enumerate(show.columns):
                        cells[j].text = str(r[c])
                doc.add_paragraph()

            doc.add_heading("Session-level results", level=1)
            add_table(sess_overall, "Table 1. Session overall (95% CI).", max_rows=12)
            doc.add_heading("Paired improvements", level=1)
            add_table(improv_overall, "Table 2. Paired improvements overall (interactive − baseline; 95% CI).", max_rows=12)

            doc.add_heading("Difficulty-stratified analysis", level=1)
            doc.add_paragraph("Difficulty values (ordered): " + (", ".join(difficulty_order) if difficulty_order else "NA"))
            add_table(sess_by_diff, "Table 3. Session by difficulty (95% CI).", max_rows=70)
            add_table(improv_by_diff, "Table 4. Paired improvements by difficulty (95% CI).", max_rows=70)

            doc.add_heading("Figures", level=1)
            for img_name, caption in [
                ("anymatch_across_sessions.png", "Any-match accuracy across sessions."),
                ("itemf1_across_sessions.png", "Diagnosis-item F1 across sessions."),
                ("time_across_sessions.png", "Mean time across sessions (seconds)."),
                ("paired_itemf1_by_difficulty_mean.png", "Mean paired improvement in item-F1 by difficulty (visual; not CI)."),
            ]:
                img_path = plots_dir / img_name
                if img_path.exists():
                    doc.add_picture(str(img_path), width=Inches(6.2))
                    cp = doc.add_paragraph(caption)
                    cp.runs[0].italic = True

            doc.save(report_docx)

    if args.make_pdf:
        rl = _try_import_reportlab()
        if rl is None:
            print("PDF requested but reportlab not available. Install: pip install reportlab")
        else:
            letter, canvas, inch, ImageReader = rl
            c = canvas.Canvas(str(report_pdf), pagesize=letter)
            W, H = letter

            def wrap_draw(text: str, x: float, y: float, width_chars: int = 105, leading: int = 12):
                c.setFont("Helvetica", 10)
                for line in textwrap.wrap(text, width=width_chars):
                    c.drawString(x, y, line)
                    y -= leading
                return y

            def draw_compact_table(df: pd.DataFrame, x: float, y: float, max_rows: int = 8):
                c.setFont("Helvetica", 8)
                cols = list(df.columns)
                header = " | ".join(cols)
                c.drawString(x, y, header[:120])
                y -= 12
                for _, r in df.head(max_rows).iterrows():
                    line = " | ".join(str(r[col]) for col in cols)
                    c.drawString(x, y, line[:120])
                    y -= 11
                return y

            y = H - 0.9 * inch
            c.setFont("Helvetica-Bold", 14)
            c.drawString(0.8 * inch, y, "MedSyn Evaluation Report (Multi-metric + Difficulty, Session-paired)")
            y -= 0.3 * inch
            y = wrap_draw(
                f"Compact PDF companion. Full tables are in the CSV exports and (optionally) DOCX. "
                f"Bootstraps={n_boot}, seed={seed}. Time treated as seconds.",
                0.8 * inch,
                y,
            )

            c.setFont("Helvetica-Bold", 12)
            y -= 8
            c.drawString(0.8 * inch, y, "Key session results (overall)"); y -= 16
            key_sess = sess_overall[["session", "condition", "any_match_acc_ci", "f1_items_ci", "mean_time_s_ci"]].copy()
            y = draw_compact_table(key_sess, 0.8 * inch, y, 6)

            y -= 10
            c.setFont("Helvetica-Bold", 12)
            c.drawString(0.8 * inch, y, "Paired improvements (overall)"); y -= 16
            key_imp = improv_overall[["pair", "any_match_acc_diff_ci", "f1_items_diff_ci", "mean_time_diff_ci"]].copy()
            y = draw_compact_table(key_imp, 0.8 * inch, y, 4)

            y -= 10
            c.setFont("Helvetica-Bold", 12)
            c.drawString(0.8 * inch, y, "Paired item-F1 by difficulty (preview)"); y -= 16
            prev = improv_by_diff[["pair", "difficulty", "f1_items_diff_ci"]].copy()
            y = draw_compact_table(prev, 0.8 * inch, y, 8)

            c.showPage()

            def draw_image(img_path: Path, caption: str, top_y: float, width_in: float = 6.8):
                if not img_path.exists():
                    return top_y
                img = ImageReader(str(img_path))
                iw, ih = img.getSize()
                tw = width_in * inch
                th = tw * (ih / iw)
                c.drawImage(img, 0.8 * inch, top_y - th, width=tw, height=th, preserveAspectRatio=True, mask="auto")
                c.setFont("Helvetica-Oblique", 9)
                c.drawString(0.8 * inch, top_y - th - 12, caption[:110])
                return top_y - th - 0.35 * inch

            top = H - 0.8 * inch
            top = draw_image(plots_dir / "itemf1_across_sessions.png", "Figure: Item-F1 across sessions.", top)
            top = draw_image(plots_dir / "paired_itemf1_by_difficulty_mean.png", "Figure: Mean paired Δ item-F1 by difficulty.", top)

            c.save()

    if args.make_tex:
        tex = []
        tex.append(r"\documentclass[11pt]{article}")
        tex.append(r"\usepackage[margin=1in]{geometry}")
        tex.append(r"\usepackage{graphicx}")
        tex.append(r"\usepackage{booktabs}")
        tex.append(r"\begin{document}")
        tex.append(r"\title{MedSyn Evaluation Report (Multi-metric + Difficulty, Session-paired)}")
        tex.append(r"\date{}")
        tex.append(r"\maketitle")

        tex.append(r"\section{Overview}")
        tex.append(
            r"Sessions are analyzed separately. LLM assistance is quantified using paired comparisons "
            r"Session~1$\rightarrow$2 and Session~3$\rightarrow$4. We stratify performance and improvements by case difficulty."
        )

        tex.append(r"\section{Key tables}")
        tex.append(df_to_latex(
            sess_overall[["session", "condition", "any_match_acc_ci", "f1_items_ci", "mean_time_s_ci"]],
            "Session-level overall results (95\\% CI).",
            "tab:sess_overall",
        ))
        tex.append(df_to_latex(
            improv_overall[["pair", "any_match_acc_diff_ci", "f1_items_diff_ci", "mean_time_diff_ci"]],
            "Paired improvements overall (interactive minus baseline; 95\\% CI).",
            "tab:paired_overall",
        ))
        tex.append(df_to_latex(
            improv_by_diff[["pair", "difficulty", "f1_items_diff_ci", "mean_time_diff_ci"]],
            "Paired improvements by difficulty (selected metrics; 95\\% CI).",
            "tab:paired_diff",
        ))

        tex.append(r"\section{Figures}")
        tex.append(r"\begin{figure}[t]\centering")
        tex.append(r"\includegraphics[width=0.95\linewidth]{plots/itemf1_across_sessions.png}")
        tex.append(r"\caption{Diagnosis-item F1 across sessions.}\end{figure}")

        if (plots_dir / "paired_itemf1_by_difficulty_mean.png").exists():
            tex.append(r"\begin{figure}[t]\centering")
            tex.append(r"\includegraphics[width=0.95\linewidth]{plots/paired_itemf1_by_difficulty_mean.png}")
            tex.append(r"\caption{Mean paired improvement in item-F1 by difficulty (visual).}\end{figure}")

        tex.append(r"\end{document}")
        report_tex.write_text("\n\n".join(tex), encoding="utf-8")

    if args.make_zip:
        with zipfile.ZipFile(report_zip, "w", zipfile.ZIP_DEFLATED) as z:
            # reports if exist
            for rp in [report_docx, report_pdf, report_tex]:
                if rp.exists():
                    z.write(rp, arcname=rp.name)

            # code itself (best-effort)
            try:
                z.write(Path(__file__), arcname=Path(__file__).name)
            except Exception:
                pass

            # tables/plots
            for f in out_dir.rglob("*"):
                if f.is_file() and f != report_zip:
                    z.write(f, arcname=str(f.relative_to(out_dir)))

    print("\nDONE.")
    print(f"Output folder: {out_dir.resolve()}")
    print(f"Long-format table: {long_path.name}")
    print(f"BERTScore computed: {bs_ok}")
    print(f"Difficulty order: {difficulty_order}")


if __name__ == "__main__":
    main()
