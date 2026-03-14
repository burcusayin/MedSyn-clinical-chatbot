#!/usr/bin/env python3
"""MedSyn manual-vs-automated evaluation (reproducible).

Inputs (defaults are the filenames you provided):
  manual_eval_session1_baseline.csv
  manual_eval_session2_interactive.csv
  manual_eval_session3_baseline.csv
  manual_eval_session4_interactive.csv

Outputs:
  manual_vs_auto_eval_outputs/
    tables/*.csv
    plots/*.png
    MedSyn_Manual_vs_Automated_Evaluation_Report.docx
    MedSyn_Manual_vs_Automated_Evaluation_Report.pdf

This script:
1) Computes manual performance (3-class rubric + binary correct vs wrong).
2) Computes automated metrics using token-based fuzzy matching (token F1 + greedy 1-to-1 matching).
3) Difficulty-standardizes per participant-session using fixed weights (Easy=3/13, Medium=6/13, Hard=4/13).
4) Evaluates RQ1–RQ3 and agreement between manual and automated labels.
"""

from __future__ import annotations
import argparse
import math
import re
import shutil
from pathlib import Path
from typing import Dict, List, Tuple

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# ---------------------------
# Configuration
# ---------------------------

PARTICIPANTS = ["phy1","phy2","phy3","res1","res3","res5","res6"]
EXPERTISE = {p: ("senior" if p.startswith("phy") else "resident") for p in PARTICIPANTS}
SESSION_COND = {1:"baseline",2:"interactive",3:"baseline",4:"interactive"}

DIFF_WEIGHTS = {"Easy":3/13, "Medium":6/13, "Hard":4/13}

TOKEN_RE = re.compile(r"[a-z0-9]+", re.I)

MAN3 = {"WRONG":0.0, "PARTIALLY CORRECT":0.5, "COMPLETELY CORRECT":1.0}
ORDER3 = ["WRONG","PARTIALLY CORRECT","COMPLETELY CORRECT"]

# ---------------------------
# Text / matching helpers
# ---------------------------

def norm_item(s: str) -> str:
    s = "" if s is None or (isinstance(s,float) and np.isnan(s)) else str(s)
    s = s.lower().strip()
    s = re.sub(r"[^\w\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def split_items(text: str) -> List[str]:
    s = "" if text is None or (isinstance(text,float) and np.isnan(text)) else str(text)
    s = s.replace("\r\n","\n").replace("\r","\n")
    parts = re.split(r"[;\n]+", s)
    out = [norm_item(p) for p in parts if norm_item(p)]
    return out

def tokset(s: str) -> set:
    return set(TOKEN_RE.findall(norm_item(s)))

def token_f1(a: str, b: str) -> float:
    A = tokset(a); B = tokset(b)
    if not A and not B: return 1.0
    if not A or not B: return 0.0
    inter = len(A & B)
    return (2*inter)/(len(A)+len(B))

def greedy_match(pred_items: List[str], gold_items: List[str], thresh: float) -> List[Tuple[int,int,float]]:
    if not pred_items or not gold_items:
        return []
    sims = []
    for i,p in enumerate(pred_items):
        for j,g in enumerate(gold_items):
            sims.append((token_f1(p,g), i, j))
    sims.sort(reverse=True, key=lambda x: x[0])
    used_p=set(); used_g=set()
    matches=[]
    for score,i,j in sims:
        if score < thresh: break
        if i in used_p or j in used_g: continue
        used_p.add(i); used_g.add(j)
        matches.append((i,j,score))
    return matches

def auto_metrics_for_case(pred_items: List[str], gold_items: List[str], thresh: float) -> Dict[str,float]:
    P = pred_items
    G = gold_items
    matches = greedy_match(P,G,thresh=thresh)
    m = len(matches)
    Pn = len(P); Gn = len(G)
    prec = m/Pn if Pn else (1.0 if Gn==0 else 0.0)
    rec  = m/Gn if Gn else (1.0 if Pn==0 else 0.0)
    f1 = 0.0 if (prec+rec)==0 else 2*prec*rec/(prec+rec)
    any_match = 1 if m>0 else 0
    exact_match = 1 if ((m==Pn==Gn and Pn>0) or (Pn==Gn==0)) else 0
    return {
        "precision":prec,"recall":rec,"f1":f1,
        "any_match":any_match,"exact_match":exact_match,
    }

def norm_manual_label(x) -> str:
    s = "" if x is None or (isinstance(x,float) and np.isnan(x)) else str(x).strip().upper()
    s = re.sub(r"\s+", " ", s)
    return s

def manual_binary(label: str) -> int:
    s = norm_manual_label(label)
    return 1 if s in ("PARTIALLY CORRECT","COMPLETELY CORRECT") else 0

# ---------------------------
# Build long-format table
# ---------------------------

def build_long(session_dfs: Dict[int,pd.DataFrame], thresh: float) -> pd.DataFrame:
    rows=[]
    for session, df in session_dfs.items():
        for idx, r in df.iterrows():
            gold_items = split_items(r["discharge diagnosis"])
            diff = r["Difficulty"]
            for p in PARTICIPANTS:
                pred_items = split_items(r.get(f"{p}_answer", ""))
                man_label = norm_manual_label(r.get(f"{p}_correctness",""))
                man_score = MAN3.get(man_label, np.nan)
                man_bin = manual_binary(man_label) if man_label else np.nan

                auto = auto_metrics_for_case(pred_items, gold_items, thresh=thresh)
                auto_3 = "WRONG"
                if auto["exact_match"]==1:
                    auto_3 = "COMPLETELY CORRECT"
                elif auto["any_match"]==1:
                    auto_3 = "PARTIALLY CORRECT"

                rows.append({
                    "session":session,
                    "condition":SESSION_COND[session],
                    "case_in_session": idx+1,
                    "difficulty": diff,
                    "participant": p,
                    "expertise": EXPERTISE[p],
                    "manual_label3": man_label,
                    "manual_score": man_score,
                    "manual_correct": man_bin,
                    "auto_precision": auto["precision"],
                    "auto_recall": auto["recall"],
                    "auto_f1": auto["f1"],
                    "auto_any_match": auto["any_match"],
                    "auto_exact_match": auto["exact_match"],
                    "auto_label3": auto_3,
                    "auto_correct": 1 if auto["any_match"]==1 else 0,
                })
    return pd.DataFrame(rows)

# ---------------------------
# Standardization / bootstrap
# ---------------------------

def standardize_by_difficulty(df: pd.DataFrame, metric_col: str) -> pd.DataFrame:
    g = df.groupby(["session","condition","participant","expertise","difficulty"])[metric_col].mean().reset_index()
    piv = g.pivot_table(index=["session","condition","participant","expertise"], columns="difficulty", values=metric_col)
    for d in DIFF_WEIGHTS:
        if d not in piv.columns:
            piv[d]=np.nan
    std = sum(DIFF_WEIGHTS[d]*piv[d] for d in DIFF_WEIGHTS)
    out = piv.reset_index()
    out[f"{metric_col}_std"] = std.values
    return out[["session","condition","participant","expertise",f"{metric_col}_std"]]

def bootstrap_ci_participants(df: pd.DataFrame, value_col: str, group_cols: List[str], n_boot: int=4000, seed: int=7) -> pd.DataFrame:
    rng = np.random.default_rng(seed)
    out=[]
    for keys, g in df.groupby(group_cols):
        participants = g["participant"].unique().tolist()
        per_p = g.groupby("participant")[value_col].mean().reindex(participants).values.astype(float)
        per_p = per_p[~np.isnan(per_p)]
        if len(per_p)==0:
            mean=lo=hi=np.nan
        else:
            mean=float(np.mean(per_p))
            boots = rng.choice(per_p, size=(n_boot,len(per_p)), replace=True).mean(axis=1)
            lo,hi = np.percentile(boots,[2.5,97.5]).astype(float)
        if not isinstance(keys, tuple):
            keys=(keys,)
        row=dict(zip(group_cols, keys))
        row["mean"]=mean; row["ci_low"]=lo; row["ci_high"]=hi
        row["ci_str"]=f"{mean:.3f} [{lo:.3f}, {hi:.3f}]"
        out.append(row)
    return pd.DataFrame(out)

def paired_improvement(std_df: pd.DataFrame, metric_std_col: str) -> pd.DataFrame:
    wide = std_df.pivot_table(index=["participant","expertise"], columns="session", values=metric_std_col)
    rows=[]
    for pair, (s_b, s_i) in {"S1->S2":(1,2), "S3->S4":(3,4)}.items():
        diff = wide[s_i] - wide[s_b]
        tmp = diff.reset_index()
        tmp["pair"]=pair
        tmp["diff"]=diff.values
        rows.append(tmp[["participant","expertise","pair","diff"]])
    return pd.concat(rows, ignore_index=True)

def bootstrap_ci_paired(paired_df: pd.DataFrame, group_cols: List[str], n_boot: int=4000, seed: int=7) -> pd.DataFrame:
    rng=np.random.default_rng(seed)
    out=[]
    for keys, g in paired_df.groupby(group_cols):
        vals = g.groupby("participant")["diff"].mean().values.astype(float)
        vals = vals[~np.isnan(vals)]
        if len(vals)==0:
            mean=lo=hi=np.nan
        else:
            mean=float(vals.mean())
            boots = rng.choice(vals, size=(n_boot,len(vals)), replace=True).mean(axis=1)
            lo,hi = np.percentile(boots,[2.5,97.5]).astype(float)
        if not isinstance(keys, tuple):
            keys=(keys,)
        row=dict(zip(group_cols, keys))
        row["diff_mean"]=mean; row["ci_low"]=lo; row["ci_high"]=hi
        row["diff_ci_str"]=f"{mean:.3f} [{lo:.3f}, {hi:.3f}]"
        out.append(row)
    return pd.DataFrame(out)

# ---------------------------
# Agreement
# ---------------------------

def confusion_matrix(labels_true: List[str], labels_pred: List[str], label_order: List[str]) -> np.ndarray:
    idx = {lab:i for i,lab in enumerate(label_order)}
    m = np.zeros((len(label_order),len(label_order)), int)
    for t,p in zip(labels_true, labels_pred):
        m[idx[t], idx[p]] += 1
    return m

def cohen_kappa_from_conf(conf: np.ndarray) -> float:
    conf = conf.astype(float)
    n = conf.sum()
    if n==0: return float("nan")
    p0 = np.trace(conf)/n
    row = conf.sum(axis=1)/n
    col = conf.sum(axis=0)/n
    pe = (row*col).sum()
    return float("nan") if pe==1 else float((p0-pe)/(1-pe))

def weighted_kappa_from_conf(conf: np.ndarray, weights: str="quadratic") -> float:
    conf=conf.astype(float)
    n=conf.sum()
    if n==0: return float("nan")
    k=conf.shape[0]
    W=np.zeros((k,k))
    for i in range(k):
        for j in range(k):
            if weights=="linear":
                W[i,j]=abs(i-j)/(k-1) if k>1 else 0
            else:
                W[i,j]=((i-j)**2)/((k-1)**2) if k>1 else 0
    O=conf/n
    row=O.sum(axis=1); col=O.sum(axis=0)
    E=np.outer(row,col)
    num=(W*O).sum()
    den=(W*E).sum()
    return float(1 - (num/den if den>0 else float("nan")))

def binary_stats(y_true: np.ndarray, y_pred: np.ndarray) -> Dict[str,float]:
    y_true=np.asarray(y_true, int); y_pred=np.asarray(y_pred, int)
    tp = int(((y_true==1)&(y_pred==1)).sum())
    tn = int(((y_true==0)&(y_pred==0)).sum())
    fp = int(((y_true==0)&(y_pred==1)).sum())
    fn = int(((y_true==1)&(y_pred==0)).sum())
    acc = (tp+tn)/(tp+tn+fp+fn) if (tp+tn+fp+fn) else float("nan")
    prec = tp/(tp+fp) if (tp+fp) else float("nan")
    rec = tp/(tp+fn) if (tp+fn) else float("nan")
    f1 = 0.0 if (not np.isfinite(prec) or not np.isfinite(rec) or (prec+rec)==0) else 2*prec*rec/(prec+rec)
    spec = tn/(tn+fp) if (tn+fp) else float("nan")
    mcc_den = math.sqrt((tp+fp)*(tp+fn)*(tn+fp)*(tn+fn)) if (tp+fp)*(tp+fn)*(tn+fp)*(tn+fn) else float("nan")
    mcc = ((tp*tn - fp*fn)/mcc_den) if np.isfinite(mcc_den) and mcc_den!=0 else float("nan")
    conf=np.array([[tn,fp],[fn,tp]])
    kappa=cohen_kappa_from_conf(conf)
    return {"tp":tp,"tn":tn,"fp":fp,"fn":fn,"acc":acc,"precision":prec,"recall":rec,"f1":f1,"specificity":spec,"mcc":mcc,"kappa":kappa}

# ---------------------------
# Plotting
# ---------------------------

def parse_ci(ci_str: str) -> Tuple[float,float,float]:
    m = re.match(r"^\s*(-?\d+(?:\.\d+)?)\s*\[\s*(-?\d+(?:\.\d+)?)\s*,\s*(-?\d+(?:\.\d+)?)\s*\]\s*$", str(ci_str).strip())
    if not m:
        return (float("nan"), float("nan"), float("nan"))
    return tuple(float(x) for x in m.groups())

def plot_metric_by_session(exp_ci_df: pd.DataFrame, metric: str, out_path: Path, title: str, ylabel: str):
    sub = exp_ci_df[exp_ci_df["metric"]==metric].copy()
    sess_order = sorted(sub["session"].unique())
    labels = [f"S{s} ({SESSION_COND[s]})" for s in sess_order]
    fig, ax = plt.subplots(figsize=(9.2,4.8))
    for exp in ["resident","senior"]:
        g = sub[sub["expertise"]==exp].set_index("session").reindex(sess_order)
        y = g["mean"].values.astype(float)
        yerr = np.vstack([y - g["ci_low"].values, g["ci_high"].values - y])
        ax.errorbar(range(len(sess_order)), y, yerr=yerr, marker="o", capsize=4, label=exp)
    ax.set_xticks(range(len(sess_order)))
    ax.set_xticklabels(labels, rotation=20, ha="right")
    ax.set_title(title)
    ax.set_xlabel("Session")
    ax.set_ylabel(ylabel)
    ax.legend()
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)

def plot_manual_distribution(long_df: pd.DataFrame, out_path: Path):
    order = ORDER3
    sub = long_df.groupby(["session","condition","expertise"])["manual_label3"].value_counts(normalize=True).unstack(fill_value=0)
    for c in order:
        if c not in sub.columns:
            sub[c]=0.0
    sub = sub[order].reset_index()
    sess_order = sorted(sub["session"].unique())
    fig, ax = plt.subplots(figsize=(10.5,4.8))
    width=0.35
    x = np.arange(len(sess_order))
    for i, exp in enumerate(["resident","senior"]):
        d = sub[sub["expertise"]==exp].set_index("session").reindex(sess_order).fillna(0)
        bottoms = np.zeros(len(sess_order))
        for cls in order:
            vals = d[cls].values.astype(float)
            ax.bar(x + (i-0.5)*width, vals, width, bottom=bottoms, label=f"{exp} - {cls}")
            bottoms += vals
    ax.set_xticks(x)
    ax.set_xticklabels([f"S{s} ({SESSION_COND[s]})" for s in sess_order], rotation=20, ha="right")
    ax.set_ylabel("Proportion of cases")
    ax.set_title("Manual evaluation distribution by session and expertise")
    ax.legend(ncol=3, fontsize=8)
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)

def plot_agreement_sensitivity(df: pd.DataFrame, out_path: Path):
    fig, ax = plt.subplots(figsize=(9.2,4.8))
    ax.plot(df["thresh"], df["bin_kappa"], marker="o", label="Binary Cohen's kappa")
    ax.plot(df["thresh"], df["tri_wkappa"], marker="o", label="3-class weighted kappa")
    ax.set_xlabel("Automated matching threshold")
    ax.set_ylabel("Agreement (kappa)")
    ax.set_title("Agreement between manual and automated labels vs threshold")
    ax.legend()
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)

# ---------------------------
# Report generation
# ---------------------------

def add_df_table_docx(doc: Document, df: pd.DataFrame, title: str, font_size: int=9):
    p = doc.add_paragraph(title)
    p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            cells[j].text = str(row[col])
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
    doc.add_paragraph("")

def build_reports(out_root: Path, tables: Dict[str,Dict[str,pd.DataFrame]], plots_dir: Path):
    # DOCX
    docx_path = out_root/"MedSyn_Manual_vs_Automated_Evaluation_Report.docx"
    doc = Document()
    t = doc.add_paragraph("MedSyn: Manual vs Automated Evaluation of Primary Diagnosis Prediction")
    t.runs[0].bold=True
    t.runs[0].font.size = Pt(16)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    doc.add_paragraph(
        "This report analyzes a clinician-led manual evaluation of diagnostic answers in four user-study sessions "
        "(S1 baseline, S2 interactive, S3 baseline, S4 interactive). We summarize manual performance (3-class rubric and binary correct/wrong) "
        "and compare manual labels with an automated fuzzy-matching evaluation to quantify alignment."
    )

    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "Manual rubric: WRONG / PARTIALLY CORRECT / COMPLETELY CORRECT. "
        "Automated evaluation: token-based similarity (token F1) + greedy one-to-one matching between predicted and gold diagnosis items; "
        "derived metrics include any-match and exact-match. "
        "Difficulty-standardization uses fixed weights Easy=3/13, Medium=6/13, Hard=4/13."
    )

    for key in ["strict_0.6","loose_0.1"]:
        doc.add_heading(f"Results (automated configuration: {key})", level=2)
        add_df_table_docx(doc, tables[key]["manual_session"], "Table A. Manual performance per session (standardized; 95% CI).")
        add_df_table_docx(doc, tables[key]["auto_session"], "Table B. Automated performance per session (standardized; 95% CI).")
        add_df_table_docx(doc, tables[key]["paired_key"], "Table C. Paired improvements (interactive − baseline; 95% CI).")
        add_df_table_docx(doc, tables[key]["paired_exp_key"], "Table D. Paired improvements by expertise (95% CI).")
        add_df_table_docx(doc, tables[key]["agreement"], "Table E. Overall agreement between manual and automated labels.")

    # Add key figures
    doc.add_heading("Figures", level=1)
    for fn in [
        "manual_distribution_session_expertise.png",
        "manual_correct_by_session_strict_0.6.png",
        "auto_anymatch_by_session_strict_0.6.png",
        "agreement_sensitivity.png",
        "scatter_manualscore_autof1_strict.png",
    ]:
        p = doc.add_paragraph(fn)
        p.runs[0].italic=True
        doc.add_picture(str(plots_dir/fn), width=Inches(6.8))
        doc.add_paragraph("")

    doc.save(docx_path)

    # PDF (landscape, wrap tables)
    pdf_path = out_root/"MedSyn_Manual_vs_Automated_Evaluation_Report.pdf"
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]
    body10 = ParagraphStyle("Body10", parent=body, fontSize=10, leading=12, spaceAfter=8)
    small9 = ParagraphStyle("Small9", parent=body, fontSize=9, leading=11, spaceAfter=6)
    tiny7 = ParagraphStyle("Tiny7", parent=body, fontSize=7, leading=9)

    PAGE = landscape(letter)
    W,H = PAGE
    left=right=0.6*inch
    usable_w = W-left-right

    def widths_for(df, first=None):
        n=len(df.columns)
        if first is None:
            return [usable_w/n]*n
        rem=usable_w-sum(first)
        k=n-len(first)
        return first + ([rem/k]*k if k>0 else [])

    def tbl_from_df(df, col_widths=None):
        data = [list(df.columns)]
        for _, r in df.iterrows():
            data.append([str(r[c]) for c in df.columns])
        pdata=[]
        for row in data:
            prow=[Paragraph(str(cell).replace("\n","<br/>"), tiny7) for cell in row]
            pdata.append(prow)
        if col_widths is None:
            col_widths=[usable_w/len(df.columns)]*len(df.columns)
        t=Table(pdata, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#E6E6E6")),
            ("GRID",(0,0),(-1,-1),0.25,colors.grey),
            ("VALIGN",(0,0),(-1,-1),"TOP"),
            ("LEFTPADDING",(0,0),(-1,-1),3),
            ("RIGHTPADDING",(0,0),(-1,-1),3),
            ("TOPPADDING",(0,0),(-1,-1),2),
            ("BOTTOMPADDING",(0,0),(-1,-1),2),
        ]))
        return t

    docpdf = SimpleDocTemplate(str(pdf_path), pagesize=PAGE, leftMargin=left, rightMargin=right,
                              topMargin=0.55*inch, bottomMargin=0.55*inch)
    story=[]
    story.append(Paragraph("MedSyn: Manual vs Automated Evaluation Report", title_style))
    story.append(Paragraph(
        "Manual labels (WRONG/PARTIALLY/COMPLETELY CORRECT) are compared against automated fuzzy matching. "
        "Intervals are 95% participant bootstrap CIs; scores are difficulty-standardized (Easy=3/13, Medium=6/13, Hard=4/13).",
        body10
    ))
    story.append(Paragraph("Results", h1))

    for key in ["strict_0.6","loose_0.1"]:
        story.append(Paragraph(f"Automated configuration: {key}", h2))
        story.append(Paragraph("Manual performance per session (standardized; 95% CI).", small9))
        t = tables[key]["manual_session"]
        story.append(tbl_from_df(t, widths_for(t, first=[0.7*inch, 1.2*inch])))
        story.append(Spacer(1,8))
        story.append(Paragraph("Automated performance per session (standardized; 95% CI).", small9))
        t = tables[key]["auto_session"]
        story.append(tbl_from_df(t, widths_for(t, first=[0.7*inch, 1.2*inch])))
        story.append(Spacer(1,8))
        story.append(Paragraph("Paired improvements (interactive − baseline; 95% CI).", small9))
        t = tables[key]["paired_key"]
        story.append(tbl_from_df(t, widths_for(t, first=[1.2*inch, 2.0*inch])))
        story.append(PageBreak())

    story.append(Paragraph("Figures", h1))
    for fn, cap in [
        ("manual_distribution_session_expertise.png","Manual label distribution by session and expertise."),
        ("manual_correct_by_session_strict_0.6.png","Manual binary correctness by session (95% CI)."),
        ("auto_anymatch_by_session_strict_0.6.png","Automated any-match by session (strict_0.6; 95% CI)."),
        ("agreement_sensitivity.png","Agreement (kappa) vs automated threshold."),
    ]:
        story.append(Paragraph(cap, small9))
        img = Image(str(plots_dir/fn))
        img.drawHeight = 3.0*inch
        img.drawWidth = 8.8*inch
        story.append(img)
        story.append(Spacer(1,10))

    docpdf.build(story)
    return docx_path, pdf_path

# ---------------------------
# Main
# ---------------------------

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--s1", default="manual_eval_session1_baseline.csv")
    ap.add_argument("--s2", default="manual_eval_session2_interactive.csv")
    ap.add_argument("--s3", default="manual_eval_session3_baseline.csv")
    ap.add_argument("--s4", default="manual_eval_session4_interactive.csv")
    ap.add_argument("--out_dir", default="manual_vs_auto_eval_outputs")
    ap.add_argument("--strict_thresh", type=float, default=0.6)
    ap.add_argument("--loose_thresh", type=float, default=0.1)
    args = ap.parse_args()

    out_root = Path(args.out_dir)
    if out_root.exists():
        shutil.rmtree(out_root)
    (out_root/"tables").mkdir(parents=True, exist_ok=True)
    (out_root/"plots").mkdir(parents=True, exist_ok=True)

    session_dfs = {
        1: pd.read_csv(args.s1),
        2: pd.read_csv(args.s2),
        3: pd.read_csv(args.s3),
        4: pd.read_csv(args.s4),
    }

    thresholds = {"strict_0.6": args.strict_thresh, "loose_0.1": args.loose_thresh}
    tables = {}

    # Build sensitivity table
    sens_rows=[]
    for th in np.linspace(0.1, 0.9, 9):
        long_df = build_long(session_dfs, float(round(th,2)))
        # binary agreement
        b = binary_stats(long_df["manual_correct"], long_df["auto_correct"])
        # 3-class agreement
        conf3 = confusion_matrix(long_df["manual_label3"].tolist(), long_df["auto_label3"].tolist(), ORDER3)
        tri_acc = np.trace(conf3)/conf3.sum()
        tri_k = cohen_kappa_from_conf(conf3)
        tri_w = weighted_kappa_from_conf(conf3, "quadratic")
        sens_rows.append({"thresh":float(round(th,2)), "bin_acc":b["acc"], "bin_kappa":b["kappa"], "tri_acc":tri_acc, "tri_kappa":tri_k, "tri_wkappa":tri_w})
    agree_sens = pd.DataFrame(sens_rows)
    agree_sens.to_csv(out_root/"tables"/"agreement_sensitivity_threshold.csv", index=False)

    for name, th in thresholds.items():
        long_df = build_long(session_dfs, th)

        # Standardized manual metrics
        man_correct_std = standardize_by_difficulty(long_df, "manual_correct")
        man_score_std = standardize_by_difficulty(long_df, "manual_score")
        man_complete_std = standardize_by_difficulty(long_df.assign(man_complete=(long_df.manual_label3=="COMPLETELY CORRECT").astype(int)), "man_complete")

        # Standardized automated metrics
        auto_any_std = standardize_by_difficulty(long_df, "auto_any_match")
        auto_exact_std = standardize_by_difficulty(long_df, "auto_exact_match")
        auto_f1_std = standardize_by_difficulty(long_df, "auto_f1")

        std = man_correct_std.merge(man_score_std, on=["session","condition","participant","expertise"])
        std = std.merge(man_complete_std, on=["session","condition","participant","expertise"])
        std = std.merge(auto_any_std, on=["session","condition","participant","expertise"])
        std = std.merge(auto_exact_std, on=["session","condition","participant","expertise"])
        std = std.merge(auto_f1_std, on=["session","condition","participant","expertise"])

        std_long = std.melt(id_vars=["session","condition","participant","expertise"], var_name="metric", value_name="value")
        sess_ci = bootstrap_ci_participants(std_long, "value", ["session","condition","metric"])
        exp_ci  = bootstrap_ci_participants(std_long, "value", ["session","condition","expertise","metric"])

        # Paired improvements
        paired_rows=[]
        for metric in [c for c in std.columns if c.endswith("_std")]:
            p_df = paired_improvement(std[["session","participant","expertise",metric]].rename(columns={metric:"metric_std"}), "metric_std")
            p_df["metric"]=metric
            paired_rows.append(p_df)
        paired_all = pd.concat(paired_rows, ignore_index=True)
        paired_ci = bootstrap_ci_paired(paired_all, ["pair","metric"])
        paired_exp_ci = bootstrap_ci_paired(paired_all, ["pair","metric","expertise"])

        # Agreement overall
        conf3 = confusion_matrix(long_df["manual_label3"].tolist(), long_df["auto_label3"].tolist(), ORDER3)
        b = binary_stats(long_df["manual_correct"], long_df["auto_correct"])
        agree = pd.DataFrame([{
            "threshold": name,
            "tri_acc": float(np.trace(conf3)/conf3.sum()),
            "tri_kappa": cohen_kappa_from_conf(conf3),
            "tri_wkappa": weighted_kappa_from_conf(conf3, "quadratic"),
            "bin_acc": b["acc"],
            "bin_kappa": b["kappa"],
            "bin_precision": b["precision"],
            "bin_recall": b["recall"],
            "bin_f1": b["f1"],
            "bin_mcc": b["mcc"],
            "tp": b["tp"], "tn": b["tn"], "fp": b["fp"], "fn": b["fn"],
        }])

        # Save tables
        sess_ci.to_csv(out_root/"tables"/f"session_ci_{name}.csv", index=False)
        exp_ci.to_csv(out_root/"tables"/f"session_by_expertise_ci_{name}.csv", index=False)
        paired_ci.to_csv(out_root/"tables"/f"paired_improvements_ci_{name}.csv", index=False)
        paired_exp_ci.to_csv(out_root/"tables"/f"paired_improvements_by_expertise_ci_{name}.csv", index=False)
        agree.to_csv(out_root/"tables"/f"agreement_overall_{name}.csv", index=False)

        # Prepare compact tables for reports
        def pivot_ci(df_ci, metric):
            sub = df_ci[df_ci.metric==metric][["session","condition","ci_str"]].copy()
            return sub.rename(columns={"ci_str":metric+" (95% CI)"})

        manual_session = pivot_ci(sess_ci, "manual_correct_std") \
            .merge(pivot_ci(sess_ci, "man_complete_std"), on=["session","condition"]) \
            .merge(pivot_ci(sess_ci, "manual_score_std"), on=["session","condition"])

        auto_session = pivot_ci(sess_ci, "auto_any_match_std") \
            .merge(pivot_ci(sess_ci, "auto_f1_std"), on=["session","condition"]) \
            .merge(pivot_ci(sess_ci, "auto_exact_match_std"), on=["session","condition"])

        metric_map = {
            "manual_correct_std":"Manual correct (binary, std)",
            "man_complete_std":"Manual complete (std)",
            "manual_score_std":"Manual ordinal score (std)",
            "auto_any_match_std":"Auto any-match (std)",
            "auto_f1_std":"Auto F1 (std)",
            "auto_exact_match_std":"Auto exact-set (std)",
        }
        key_metrics = list(metric_map.keys())
        paired_key = paired_ci[paired_ci.metric.isin(key_metrics)][["pair","metric","diff_ci_str"]].copy()
        paired_key["metric"] = paired_key["metric"].map(metric_map)

        paired_exp_key = paired_exp_ci[paired_exp_ci.metric.isin(["manual_correct_std","man_complete_std","auto_any_match_std","auto_f1_std"])][["pair","expertise","metric","diff_ci_str"]].copy()
        paired_exp_key["metric"] = paired_exp_key["metric"].map(metric_map)

        tables[name] = {
            "manual_session": manual_session,
            "auto_session": auto_session,
            "paired_key": paired_key,
            "paired_exp_key": paired_exp_key,
            "agreement": agree,
        }

        # Plots
        plot_metric_by_session(exp_ci, "manual_correct_std", out_root/"plots"/f"manual_correct_by_session_{name}.png",
                               f"Manual correctness by session — {name}", "Standardized correct rate")
        plot_metric_by_session(exp_ci, "man_complete_std", out_root/"plots"/f"manual_complete_by_session_{name}.png",
                               f"Manual complete by session — {name}", "Standardized complete rate")
        plot_metric_by_session(exp_ci, "auto_any_match_std", out_root/"plots"/f"auto_anymatch_by_session_{name}.png",
                               f"Automated any-match by session — {name}", "Standardized any-match")
        plot_metric_by_session(exp_ci, "auto_f1_std", out_root/"plots"/f"auto_f1_by_session_{name}.png",
                               f"Automated F1 by session — {name}", "Standardized F1")

    # Common plots
    # Use strict run for manual distribution plot (manual labels do not depend on threshold)
    strict_long = build_long(session_dfs, thresholds["strict_0.6"])
    plot_manual_distribution(strict_long, out_root/"plots"/"manual_distribution_session_expertise.png")
    plot_agreement_sensitivity(agree_sens, out_root/"plots"/"agreement_sensitivity.png")

    # Reports
    docx_path, pdf_path = build_reports(out_root, tables, out_root/"plots")

    print("DONE")
    print("Output folder:", out_root.resolve())
    print("DOCX:", docx_path.resolve())
    print("PDF :", pdf_path.resolve())

if __name__ == "__main__":
    main()
